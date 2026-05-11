import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

base_dir = r"C:\Users\PC\.gemini\antigravity\scratch\korea_logistics_hub_project"

def generate_report():
    document = Document()
    
    # Title
    title = document.add_heading('Báo Cáo Nghiên Cứu: Tối Ưu Vị Trí Trung Tâm Logistics', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = document.add_paragraph('LogiHub Intelligence Core Engine - 2023 Base Year')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Section 1
    document.add_heading('1. Tổng Quan Dữ Liệu (Data Preparation)', level=1)
    p = document.add_paragraph()
    p.add_run('Dự án đã tích hợp và xử lý thành công hai nguồn dữ liệu lớn để làm đầu vào cho mô hình tối ưu:\n').bold = True
    document.add_paragraph('• Dữ liệu Nhu cầu (Demand): Dựa trên ma trận vận chuyển hàng hóa KTDB Freight O/D 2023. Gyeonggi-do, Seoul và Incheon chiếm tỷ trọng giao thương lớn nhất.', style='List Bullet')
    document.add_paragraph('• Dữ liệu Năng lực (Capacity): Cập nhật từ tệp đăng ký kho bãi thực tế (물류창고정보). Gyeonggi-do dẫn đầu với hơn 2,200 kho bãi và diện tích lên đến 24.5 triệu m².', style='List Bullet')
    document.add_paragraph('• Ma trận khoảng cách (Distance Matrix): Áp dụng khoảng cách Geodesic dựa trên tọa độ trung tâm của 17 vùng hành chính.', style='List Bullet')
    
    # Add pictures if they exist
    fig_dir = os.path.join(base_dir, "outputs", "figures")
    pic1 = os.path.join(fig_dir, "total_volume_by_region.png")
    pic2 = os.path.join(fig_dir, "top_corridors.png")
    
    if os.path.exists(pic1):
        document.add_picture(pic1, width=Inches(6.0))
        document.add_paragraph('Biểu đồ 1: Tổng khối lượng hàng hóa theo 17 vùng hành chính (2023)').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Section 2
    document.add_heading('2. Kết Quả Mô Hình P-Median (RQ3)', level=1)
    p = document.add_paragraph('Mô hình P-Median được chạy nhằm xác định vị trí Hub sao cho tổng khoảng cách vận chuyển (được nhân với trọng số lượng hàng hóa) là nhỏ nhất. Kết quả thay đổi số lượng Hub (P) mang lại các Insight sau:')
    
    document.add_paragraph('• Kịch bản 3 Hubs (P=3): Tổng chi phí vận tải đạt 162.3 tỷ Tấn-km. Các Hub được chọn: Gyeonggi-do, Jeollanam-do, Busan.', style='List Bullet')
    document.add_paragraph('• Kịch bản 5 Hubs (P=5): Tổng chi phí vận tải giảm mạnh xuống 93.7 tỷ Tấn-km (giảm 42%). Các Hub bổ sung: Chungcheongnam-do, Gyeongsangbuk-do.', style='List Bullet')
    document.add_paragraph('• Kịch bản 7 Hubs (P=7): Tổng chi phí vận tải đạt 58.2 tỷ Tấn-km.', style='List Bullet')
    
    document.add_paragraph('Kết luận: Việc mở rộng từ 3 lên 5 Hubs đem lại hiệu quả biên (Marginal Return) rất lớn. Tuy nhiên, nếu mở quá 5 Hubs, tốc độ giảm chi phí sẽ chậm lại, dẫn đến tình trạng dư thừa năng lực đầu tư so với hiệu quả tiết kiệm chi phí.', style='Intense Quote')
    
    # Section 3
    document.add_heading('3. Kết Quả Mô Hình UFLP & CFLP (RQ4)', level=1)
    document.add_paragraph('Mô hình Capacitated Facility Location Problem (CFLP) đánh giá mạng lưới với Ràng buộc sức chứa kho bãi thực tế và Chi phí cố định để vận hành Hub.')
    document.add_paragraph('• Kết quả Baseline: Để xử lý toàn bộ nhu cầu luân chuyển hàng hóa năm 2023 của Hàn Quốc với chi phí tối thiểu, mạng lưới cần kích hoạt 11 Hub trung tâm.', style='List Bullet')
    document.add_paragraph('• Danh sách 11 Hub được chọn: Gyeonggi-do, Chungcheongnam-do, Jeollanam-do, Busan, Ulsan, Incheon, Gyeongsangbuk-do, Gyeongsangnam-do, Chungcheongbuk-do, Jeollabuk-do, Sejong.', style='List Bullet')
    
    # Section 4
    document.add_heading('4. Phân Tích Kịch Bản & Độ Ổn Định (RQ5-RQ7)', level=1)
    document.add_paragraph('Để đánh giá tính bền bỉ của mạng lưới (Network Robustness), một kịch bản "Cú sốc nhu cầu" (Demand Shock) đã được giả lập bằng cách tăng tổng sản lượng hàng hóa thêm 20%.')
    document.add_paragraph('• Tác động Chi phí: Tổng chi phí toàn mạng lưới tăng từ 89.2 tỷ lên 97.5 tỷ (+9.3%). Mạng lưới thể hiện khả năng chịu tải tốt.', style='List Bullet')
    document.add_paragraph('• Thay đổi Cấu trúc: Hệ thống vẫn duy trì thiết kế 11 Hub, nhưng đã tự động hoán đổi vị trí của Sejong (đã đạt giới hạn sức chứa) bằng Daegu (để tận dụng công suất nhàn rỗi ở khu vực lân cận).', style='List Bullet')
    document.add_paragraph('Chiến lược đề xuất: Đối với bài toán quy hoạch dài hạn (2025-2050), Daegu nên được đưa vào danh mục dự án chiến lược (Strategic Pipeline) như một Hub dự phòng (Overflow Hub) để giảm tải cho khu vực miền Trung và miền Nam khi sản lượng tăng vọt.', style='Intense Quote')
    
    # Final Output
    out_dir = os.path.join(base_dir, "reports")
    out_file = os.path.join(out_dir, "LogiHub_Optimization_Report.docx")
    document.save(out_file)
    print(f"Document saved to: {out_file}")

if __name__ == "__main__":
    generate_report()
